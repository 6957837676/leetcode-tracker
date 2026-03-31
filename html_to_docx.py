#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简历生成工具 - 余荣华
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_resume():
    doc = Document()
    
    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.15
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('余荣华')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 求职意向
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('求职意向：测试工程师丨测试开发工程师')
    run.font.size = Pt(11)
    run.font.bold = True
    
    # 联系方式
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('联系电话：17858506683  |  邮箱：695783676@qq.com')
    run.font.size = Pt(10)
    
    # 核心标签
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run('核心标签：6年测试经验|测试开发|AI测试工程化落地|全流程质量保障|自动化工具研发')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(26, 54, 93)
    
    # 一、教育背景
    add_section(doc, '一、教育背景')
    add_edu(doc, '2018.09 - 2020.06', '哈尔滨工程大学', '机电一体化', '硕士研究生')
    add_edu(doc, '2014.09 - 2018.06', '绍兴理工学院', '机械设计及其自动化', '本科 工程学士')
    
    # 二、工作经历
    add_section(doc, '二、工作经历')
    
    # Shopee
    add_job(doc, '虾皮科技有限公司', '测试开发工程师', '2022.05 - 至今', '深圳')
    add_desc(doc, '跨境消费贷业务测试负责人，主导核心交易模块全生命周期质量保障，参与测试体系搭建与AI提效工具落地，支撑东南亚8国业务迭代。')
    
    add_subsection(doc, '业务质量保障体系建设')
    add_bullet(doc, '主导跨境消费贷交易核心模块全生命周期质量保障，覆盖退款、减免等核心功能30+服务接口，落地2000+全量测试用例，支撑东南亚8个国家及地区业务落地，实现核心业务上线零重大线上故障。')
    add_bullet(doc, '全流程管控版本交付质量，牵头需求评审、影响范围评估、核心场景回归策略制定，搭建需求准入标准与上线流程规范，保障需求交付质量与版本上线稳定性。')
    add_bullet(doc, '负责线上故障应急响应与根因分析，快速定位问题并推动研发侧整改落地，通过故障复盘优化测试流程与用例覆盖，持续提升系统健壮性。')
    add_bullet(doc, '参与系统容灾演练方案设计与落地验证，保障灾难恢复流程有效性与环境切换平滑性，确保跨境支付信贷业务7×24小时连续稳定运行。')
    
    add_subsection(doc, '自动化测试体系搭建')
    add_bullet(doc, '基于Python+Behave(BDD)搭建并维护接口自动化测试体系，实现核心业务接口100%自动化覆盖，持续优化自动化用例和流程，大幅提升回归测试效率。')
    add_bullet(doc, '精通接口全流程测试，熟练使用Postman、Requests完成接口用例设计与调试，通过WireShark抓包完成问题定位，结合MySQL和Trace链路跟踪平台，完成全链路数据一致性校验，保障金融场景资金与数据零差错。')
    
    add_subsection(doc, 'AI测试工程化与提效落地')
    add_bullet(doc, '从0到1主导一站式QA智能测试平台的架构设计、全模块开发与落地推广，覆盖测试用例智能设计、自动化问题排查等核心测试环节；同步自研测试全流程节点提醒工具，补齐流程管控能力，解决业务迭代快、人工成本高、流程管控难等核心痛点。')
    add_bullet(doc, '深耕LLM在测试领域的工程化落地，熟悉Prompt工程、Function Calling、Agent设计、Skills开发等核心技术，落地用例自动生成、测试数据智能构造等核心功能，通过AI技术实现测试全流程提效降本。')
    
    # 中兴
    add_job(doc, '中兴通讯有限公司', '软件测试工程师', '2020.07 - 2022.05', '深圳')
    add_bullet(doc, '负责基站测试平台的环境搭建与维护，落地环境自动化部署能力，将单套环境搭建耗时从15min缩短至5min，保障测试活动高效开展。')
    add_bullet(doc, '负责核心业务功能与接口测试，执行测试用例，精准定位缺陷并输出标准化缺陷报告，推动缺陷闭环修复，保障产品交付质量。')
    add_bullet(doc, '参与自动化测试脚本设计与编写，部署维护自动化测试环境，实现核心回归场景自动化覆盖，提升测试执行效率。')
    
    # 三、核心项目经历
    add_section(doc, '三、核心项目经历')
    
    # 项目1
    add_project(doc, '项目1：泰国SPL接入ShopeePay APP项目')
    add_project_desc(doc, '项目描述：ShopeePay东南亚业务扩张中，泰国地区存在大量纯ShopeePay账号用户，而现有信贷系统仅适配Shopee主账号体系，无法识别纯ShopeePay账号，导致消费贷交易链路断裂，严重阻碍泰国地区业务落地。本次改造实现双账号体系兼容适配，通过内部UID+业务类型的交互模式，完成业务逻辑与账号体系解耦，保障交易链路完整可用。')
    add_role(doc, '个人职责：项目测试负责人，全流程主导测试方案设计、用例落地与上线质量管控。')
    add_bullet(doc, '制定流量驱动的测试覆盖方案，拉取线上高频接口流量为基准，结合核心业务场景分析自动化覆盖缺口，避免线上场景遗漏。')
    add_bullet(doc, '采用前后端分层验证策略，前端聚焦各端APP核心场景的账号交互正确性，后端校验多账号类型下跨服务交互的账号一致性，同时覆盖8个国家及地区的差异化业务规则。')
    add_bullet(doc, '通过参数化用例设计，实现一套用例兼容多账号体系全场景验证；借助分布式追踪工具完成服务外调账号信息全链路校验，结合增量代码覆盖率分析，确保测试无死角。')
    add_bullet(doc, '开发接口请求日志采集工具，实现接口请求自动记录与覆盖度统计，保障接口覆盖完整性与调用正确性。')
    add_result(doc, '覆盖消费贷全生命周期150+服务接口，落地4000+测试用例，分3个阶段完成上线，实现泰国地区双账号体系业务顺利落地，上线零重大线上故障；沉淀多账号体系兼容性测试方法论，大幅降低同类业务测试成本。')
    
    # 项目2
    add_project(doc, '项目2：SPL支付多次退款项目')
    add_project_desc(doc, '项目描述：SPL支付业务中分笔退款需求激增，原有系统仅支持单次全额退款，无法满足用户与业务侧的分笔退款诉求，亟需完成系统能力升级。本次改造实现退款接口幂等性设计与完善的重试机制，全链路保障多次退款场景下的资金安全与数据准确性。')
    add_role(doc, '个人职责：核心测试负责人，负责全场景测试方案设计、专项测试落地与资金安全校验。')
    add_bullet(doc, '梳理订单全生命周期状态矩阵，设计覆盖「支付N次分笔退款订单完成/关闭」全链路的状态流转测试用例，全面覆盖正常与异常边界场景。')
    add_bullet(doc, '针对用户端、客服平台多入口设计交叉退款场景，验证多渠道操作下的订单状态同步性与资金数据一致性。')
    add_bullet(doc, '构建幂等性与异常重试专项测试场景，通过重复提交请求验证接口幂等性，模拟服务中断、超时等异常场景触发重试机制，校验全链路数据一致性。')
    add_bullet(doc, '执行前后端分层资金校验，前端验证退款金额展示、操作边界限制的准确性，后端通过接口断言与数据库校验，确保每笔退款的资金扣减、流水记录、数据上报的精准性。')
    add_result(doc, '实现多地区、多币种分笔退款业务全量上线，覆盖全场景边界条件，上线后零资损事件、零客诉问题；沉淀金融支付场景幂等性测试、资金安全校验体系，复用至其他交易核心模块，保障跨境支付业务资金安全。')
    
    # 项目3
    add_project(doc, '项目3：QA智能测试平台（全流程AI测试提效套件）')
    add_project_desc(doc, '项目描述：跨境消费贷业务迭代速度快，测试团队面临全流程四大核心痛点：一是自动化用例失败后根因排查耗时长，单次排查需5-20分钟；二是测试用例设计高度依赖人工，单需求耗时0.5-2天；三是测试数据构造依赖人工，单次造数耗时10-20分钟，人力占用高、易出错；四是测试流程跨环节多、节点易逾期、人工跟进成本高。本次项目搭建一站式AI测试提效平台，覆盖测试全流程核心环节，通过LLM技术+自动化流程管控，解决各环节效率与质量痛点，降低人工成本，支撑东南亚跨境业务快速迭代。')
    add_role(doc, '个人职责：项目负责人，独立完成平台整体架构设计、全模块开发、落地推广与持续迭代优化。')
    add_bullet(doc, '设计并落地智能测试用例设计模块，实现多格式需求文档解析、测试点自动提取、用例自动生成与质量管控，适配跨境业务特性，实现中文用例一键多语言翻译，对接MeterSphere平台实现测试计划自动创建。')
    add_bullet(doc, '搭建智能测试数据构造模块，开发自然语言驱动的用例智能推荐模块，基于Prompt工程搭建BDD用例智能解析系统，设计三层地区过滤机制，实现临时用例动态生成与参数注入，完成自助化测试数据智能构造。')
    add_bullet(doc, '开发自动化问题排查模块，基于LLM实现用例失败日志智能解析、根因快速定位与解决方案推荐，提升问题解决效率、赋能智能测试数据构造模块。')
    add_bullet(doc, '内置测试全流程节点管控提醒模块，对接Jira与企业微信，实现测试节点自定义配置、进度实时监控与待办/逾期消息自动推送，补齐流程自动化管控能力。')
    add_bullet(doc, '技术栈覆盖Python、FastAPI、Flask、LLM、ChromaDB、Git Worktree、BDD、WebSocket，保障平台高可用、可扩展。')
    add_result(doc, '智能用例设计模块，单需求用例设计耗时从4-8小时缩短至0.5-2小时，用例准确率达90%以上；智能造数模块，单次造数耗时从10-20分钟缩短至1-2分钟，节省测试人工支持人力30%，造数执行成功率达95%以上；自动化问题排查模块，单故障根因定位耗时从5-20分钟缩短至1-3分钟，排查准确率达85%以上；流程管控模块有效降低测试节点逾期率，减少人工跟进成本，提升跨团队协作效率。平台已在团队内全量推广落地，覆盖所有跨境消费贷业务测试项目，成为团队测试提效核心工具。')
    
    # 四、专业技能
    add_section(doc, '四、专业技能')
    add_bullet(doc, '质量保障与测试管理：6年金融信贷、通信基站业务测试实战经验，精通全流程测试方案、用例设计、风险管控，具备丰富的资损防控、幂等性校验、跨境多区域适配测试经验，可独立负责核心业务全生命周期质量保障。')
    add_bullet(doc, '测试开发与自动化：精通Python开发，熟练使用Behave(BDD)、RobotFramework搭建自动化测试体系，精通接口测试全流程与MySQL数据校验，可使用FastAPI搭建后端服务，熟练使用Jira、MeterSphere等测试管理工具。')
    add_bullet(doc, 'AI测试工程化：具备从0到1落地AI测试提效工具全流程经验，精通LLM在测试场景的工程化应用，熟练掌握Prompt工程、Function Calling、Agent设计，覆盖智能造数、自动化问题排查、用例自动生成等核心场景，熟练使用Claude API完成生产级工具开发，了解LangChain框架。')
    
    # 五、个人评价
    add_section(doc, '五、个人评价')
    add_bullet(doc, '【质量保障】6年ToB/ToC核心业务测试实战经验，深度覆盖金融信贷、通信两大领域，精通服务端、接口、全链路全流程测试，具备极强的质量管控意识与风险预判能力，主导的核心业务上线零重大故障、零资损事件。')
    add_bullet(doc, '【工程能力】精通Python开发，熟练搭建自动化测试体系与后端服务，具备全栈开发基础与成熟的工程化落地能力，可独立完成测试提效工具、平台类产品从0到1的设计、开发与推广。')
    add_bullet(doc, '【AI提效】深耕LLM在测试领域的工程化应用，具备AI测试平台全流程落地经验，可精准拆解业务痛点并通过大模型技术实现测试全流程提效降本，大幅降低团队人工成本。')
    add_bullet(doc, '【团队协作】具备丰富的跨团队、跨区域协作经验，可高效联动产品、研发、运维等多角色推动问题闭环与流程优化；持续深耕测试领域前沿技术，沉淀可复用的测试方法论与工具，赋能团队整体能力提升。')
    
    # 保存
    doc.save(r'd:\leetcode\余荣华-测试开发工程师简历.docx')
    print('简历已生成：d:\\leetcode\\余荣华-测试开发工程师简历.docx')

def add_section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(26, 54, 93)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_edu(doc, date, school, major, degree):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{date}  {school}  {major}  {degree}')
    run.font.size = Pt(10)

def add_job(doc, company, position, date, location):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run1 = p.add_run(f'{company}  ')
    run1.font.bold = True
    run1.font.size = Pt(11)
    run2 = p.add_run(f'{position} | {date} | {location}')
    run2.font.size = Pt(10)

def add_desc(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(74, 85, 104)

def add_subsection(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(26, 54, 93)

def add_project(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(26, 54, 93)

def add_project_desc(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(74, 85, 104)

def add_role(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = True

def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('• ' + text)
    run.font.size = Pt(10)

def add_result(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('项目成果：' + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(5, 150, 105)

if __name__ == '__main__':
    create_resume()
